import streamlit as st
import requests
import PyPDF2
from docx import Document
import io
import re
import PyPDF2
from PIL import Image
import pytesseract
# ---------------------------
# Load API Key securely
# ---------------------------
from dotenv import load_dotenv
import os

load_dotenv()  # load local .env for development

api_key = os.getenv("GROQ_API_KEY")  # Render will use environment variable

if not api_key:
    st.error("⚠️ API key missing! Please set GROQ_API_KEY in Render environment variables.")
    st.stop()



# ✅ Configurable Tesseract OCR path with fallbacks
# Priority: ENV var TESSERACT_CMD -> Windows default path -> system PATH
_tesseract_env = os.getenv("TESSERACT_CMD")
_windows_default = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"
try:
    if _tesseract_env and os.path.exists(_tesseract_env):
        pytesseract.pytesseract.tesseract_cmd = _tesseract_env
    elif os.name == "nt" and os.path.exists(_windows_default):
        pytesseract.pytesseract.tesseract_cmd = _windows_default
    # else: rely on PATH; if not present, OCR calls will raise which we handle
except Exception:
    # If configuration fails, we let _process_image handle exceptions gracefully
    pass


class DocumentProcessor:
    def __init__(self):
        self.supported_formats = {
            ".txt", ".pdf", ".docx", ".doc", ".png", ".jpg", ".jpeg"
        }

    def process_document(self, filepath):
        """
        Detect file type and extract text safely.
        Always returns a string (even if it's an error).
        """
        try:
            file_extension = os.path.splitext(filepath)[1].lower()

            if file_extension == ".txt":
                return self._process_txt(filepath)
            elif file_extension == ".pdf":
                return self._process_pdf(filepath)
            elif file_extension in [".docx", ".doc"]:
                return self._process_word(filepath)
            elif file_extension in [".png", ".jpg", ".jpeg"]:
                return self._process_image(filepath)
            else:
                return f"[Unsupported file format: {file_extension}]"

        except Exception as e:
            return f"[File processing error: {str(e)}]"

    # ------------------------
    # File Type Processors
    # ------------------------

    def _process_txt(self, filepath):
        """Extract text from .txt files."""
        try:
            with open(filepath, "r", encoding="utf-8") as file:
                content = file.read()
            return self._clean_text(content)
        except UnicodeDecodeError:
            try:
                with open(filepath, "r", encoding="latin-1") as file:
                    content = file.read()
                return self._clean_text(content)
            except Exception as e:
                return f"[Error reading TXT file: {str(e)}]"

    def _process_pdf(self, filepath):
        """Extract text from PDF files."""
        try:
            with open(filepath, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                content = ""
                for page in pdf_reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        content += extracted + "\n"
            return self._clean_text(content) if content else "[No text extracted from PDF]"
        except Exception as e:
            return f"[Error reading PDF: {str(e)}]"

    def _process_word(self, filepath):
        """Extract text from Word documents (.docx and .doc)."""
        try:
            doc = Document(filepath)
            content = ""

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content += cell.text + " "
                content += "\n"

            return self._clean_text(content) if content else "[No text extracted from Word file]"
        except Exception as e:
            return f"[Error reading Word document: {str(e)}]"

    def _process_image(self, filepath):
        """Extract text from image files using OCR (safe with timeout)."""
        try:
            with Image.open(filepath) as img:
                img = img.convert("RGB")  # ensure format is consistent
                # Add timeout to prevent hanging on large images
                content = pytesseract.image_to_string(img, timeout=30)
                cleaned = self._clean_text(content)
                return cleaned if cleaned else "[No text detected in image]"
        except Exception as e:
            # Be explicit when Tesseract is missing to help users
            if "tesseract is not installed" in str(e).lower() or "not found" in str(e).lower():
                return "[OCR unavailable: Tesseract not found. Install Tesseract or set TESSERACT_CMD]"
            return f"[Error reading image: {str(e)}]"

    # ------------------------
    # Helpers
    # ------------------------

    def _clean_text(self, text):
        """Clean and normalize extracted text."""
        if not text:
            return ""

        # Remove extra whitespace
        text = re.sub(r"\s+", " ", text)

        # Remove unwanted characters but keep punctuation
        text = re.sub(r"[^\w\s\.\,\!\?\;\:\-\(\)\[\]\{\}]", "", text)

        # Normalize spaces
        text = re.sub(r"\s+", " ", text)

        return text.strip()

    def get_document_summary(self, content, max_length=500):
        """Generate a brief summary of the document content."""
        if not content:
            return "[No content available to summarize]"

        if len(content) <= max_length:
            return content

        # Take first few sentences
        sentences = content.split(".")
        summary = ""
        for sentence in sentences:
            if len(summary + sentence) < max_length:
                summary += sentence.strip() + ". "
            else:
                break

        return summary.strip()


doc_processor = DocumentProcessor()

# ---------------------------
# Page Config
# ---------------------------
st.set_page_config(page_title="LectureBuddies - AI Chatbot", page_icon="🤖", layout="wide", initial_sidebar_state="expanded")

# ---------------------------
# Load API Key from .env
# ---------------------------
load_dotenv()
api_key = os.getenv("GROQ_API_KEY")  # make sure your .env has GROQ_API_KEY=your_api_key_here

if not api_key:
    st.error("⚠️ API key missing! Please check your .env file.")
    st.stop()

# ---------------------------
# Session Initialization
# ---------------------------
def init_session_state():
    defaults = {
        "messages": [],
        "uploaded_files": [],
        "document_contents": {}
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_session_state()

# ---------------------------
# Inject File Content Helper
# ---------------------------
def inject_file_content(user_message: str) -> str:
    """
    Replace file references in user message with extracted text,
    so model never says 'I can't see images'.
    """
    for fname, content in st.session_state.document_contents.items():
        if fname.lower() in user_message.lower():
            extracted = content if content.strip() else "[No text extracted from this file]"
            user_message = user_message.replace(
                fname,
                f"(Extracted content: {extracted[:1000]}...)"
            )
    return user_message

# ---------------------------
# API Interaction
# ---------------------------
def get_groq_response(user_input, model="llama-3.1-8b-instant"):

    """Send query + context to Groq API and return assistant response."""
    if not api_key or api_key.strip() == "":
        return "⚠️ Missing API key. Please set GROQ_API_KEY in your .env file."

    url = "https://api.groq.com/openai/v1/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}

    # Build document context
    doc_context = ""
    if st.session_state.document_contents:
        doc_context = "\n\n**Available Documents:**\n"
        for fname, content in st.session_state.document_contents.items():
            snippet = content[:1000] if content else "[No extractable text]"
            doc_context += f"\n--- {fname} ---\n{snippet}...\n"

    # Build conversation
    messages = [{"role": m["role"], "content": m["content"]} for m in st.session_state.messages]
    system_msg = (
        f"You are LectureBuddies, an AI chatbot designed for education. "
        f"Answer clearly, summarize effectively, and explain concepts step by step."
        f"{' Available Documents: ' + doc_context if doc_context else ''}\n\n"
        "Guidelines:\n"
        "📚 Education-focused\n"
        "📝 Summarization expert\n"
        "🎯 Clarity first (simple language, then details)\n"
        "✅ Confidence + accuracy\n"
        "Break down topics step-by-step, use examples, and stay professional yet supportive."
    )
    messages.insert(0, {"role": "system", "content": system_msg})
    messages.append({"role": "user", "content": user_input})

    payload = {"model": model, "messages": messages, "temperature": 0.7, "max_tokens": 1000}

    try:
        resp = requests.post(url, headers=headers, json=payload, timeout=30)
        if resp.status_code == 200:
            data = resp.json()
            return data.get("choices", [{}])[0].get("message", {}).get("content", "⚠️ No response received.")
        elif resp.status_code == 401:
            return "❌ Invalid API key. Please check GROQ_API_KEY in your .env file."
        elif resp.status_code == 429:
            return "⏳ Too many requests. Please slow down and retry shortly."
        else:
            return f"⚠️ API Error {resp.status_code}: {resp.text}"
    except requests.exceptions.Timeout:
        return "⏳ Request timed out. Please retry."
    except requests.exceptions.RequestException as e:
        return f"🌐 Network error: {e}"
    except Exception as e:
        return f"⚠️ Unexpected error: {e}"

# ---------------------------
# Document Processing
# ---------------------------
def process_document(uploaded_file):
    """Extract text from uploaded documents (txt, pdf, docx, images with OCR)."""
    try:
        # Save the uploaded file temporarily
        temp_path = os.path.join("temp", uploaded_file.name)
        os.makedirs("temp", exist_ok=True)
        with open(temp_path, "wb") as f:
            f.write(uploaded_file.read())

        # Use your DocumentProcessor
        content = doc_processor.process_document(temp_path)

        return content if content.strip() else "[No text extracted]"
    except Exception as e:
        return f"[File processing error: {e}]"

# ---------------------------
# Enhanced Styling (Matching Quiz Generator Theme)
# ---------------------------
st.markdown(
    """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;600;700&display=swap');
    
    .main-title {
        text-align: center;
        font-size: 40px;
        font-weight: 800;
        background: linear-gradient(90deg, #4e54c8, #8f94fb, #4e54c8);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        margin-bottom: -5px;
        font-family: 'Poppins', sans-serif;
    }
    .tagline {
        text-align: center;
        font-size: 16px;
        color: #666;
        margin-bottom: 15px;
        font-family: 'Poppins', sans-serif;
    }
    hr.gradient {
        border: none;
        height: 3px;
        background: linear-gradient(90deg, #4e54c8, #8f94fb, #4e54c8);
        border-radius: 50px;
        margin: 15px 0;
    }
    .custom-btn {
        display: inline-flex;
        align-items: center;
        gap: 6px;
        padding: 10px 24px;
        font-size: 14px;
        font-weight: 700;
        color: white;
        background: linear-gradient(90deg, #4e54c8, #8f94fb);
        border: none;
        border-radius: 25px;
        cursor: pointer;
        transition: all 0.3s ease;
        text-align: center;
        font-family: 'Poppins', sans-serif;
        box-shadow: 0 3px 10px rgba(78, 84, 200, 0.3);
        margin: 5px 0;
    }
    .custom-btn:hover {
        background: linear-gradient(90deg, #8f94fb, #4e54c8);
        transform: scale(1.03) translateY(-1px);
        box-shadow: 0 5px 15px rgba(78, 84, 200, 0.4);
    }
    .clear-btn {
        background: linear-gradient(90deg, #ff6b6b, #ee5a52);
        color: white;
        border: none;
        border-radius: 15px;
        padding: 8px 16px;
        font-weight: 600;
        cursor: pointer;
        transition: all 0.3s ease;
        font-size: 14px;
        margin: 5px 0;
    }
    .clear-btn:hover {
        background: linear-gradient(90deg, #ee5a52, #ff6b6b);
        transform: scale(1.03);
    }
    .chat-header {
        color: #4e54c8;
        font-size: 22px;
        font-weight: 700;
        margin-bottom: 10px;
        text-align: center;
        font-family: 'Poppins', sans-serif;
    }
    .user-message {
        text-align: right;
        margin: 6px 0;
    }
    .user-bubble {
        display: inline-block;
        background: linear-gradient(90deg, #4e54c8, #8f94fb);
        color: white;
        padding: 10px 14px;
        border-radius: 18px 18px 4px 18px;
        max-width: 70%;
        font-family: 'Poppins', sans-serif;
        font-size: 14px;
        box-shadow: 0 2px 8px rgba(78, 84, 200, 0.3);
    }
    .assistant-message {
        margin: 6px 0;
        color: #202123;
    }
    .assistant-bubble {
        background: white;
        padding: 10px 14px;
        border-radius: 18px 18px 18px 4px;
        max-width: 70%;
        border: 1px solid #e0e0e0;
        font-family: 'Poppins', sans-serif;
        font-size: 14px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.05);
    }
    .welcome-container {
        text-align: center;
        padding: 20px;
        background: linear-gradient(135deg, #f8f9ff, #e8ecff);
        border-radius: 15px;
        border: 1px solid #d1d9ff;
        margin: 10px 0;
    }
    .welcome-title {
        color: #4e54c8;
        font-size: 20px;
        font-weight: 600;
        margin-bottom: 10px;
        font-family: 'Poppins', sans-serif;
    }
    .welcome-text {
        color: #666;
        font-size: 14px;
        font-family: 'Poppins', sans-serif;
    }
    /* Professional chat input styling */
    .stChatInput > div > div > input {
        border: 2px solid #e0e0e0;
        border-radius: 25px;
        padding: 12px 16px;
        font-family: 'Poppins', sans-serif;
        font-size: 14px;
        background: white;
        box-shadow: 0 2px 8px rgba(0,0,0,0.05);
        transition: all 0.3s ease;
    }
    .stChatInput > div > div > input:focus {
        border-color: #4e54c8;
        box-shadow: 0 4px 12px rgba(78, 84, 200, 0.2);
        outline: none;
    }
    /* Fixed page height - prevent scrolling */
    .stApp {
        max-height: 100vh;
        overflow: hidden;
    }
    .block-container {
        padding-top: 1rem;
        padding-bottom: 0rem;
        padding-left: 1rem;
        padding-right: 1rem;
        max-height: 90vh;
        overflow-y: auto;
    }
    /* Reduce sidebar padding */
    .css-1d391kg {
        padding: 0.5rem;
    }
    /* Style sidebar elements */
    .stSidebar > div > div {
        font-family: 'Poppins', sans-serif;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# ---------------------------
# Sidebar (Enhanced to Match Theme)
# ---------------------------
with st.sidebar:
    st.markdown("## ⚙️ Chat Settings")
    st.markdown("### Customize your experience")
    
    if st.button("🗑️ Clear Chat", key="clear_chat", help="Start a new conversation"):
        st.session_state.messages.clear()
        st.rerun()

    st.markdown("---")
    st.markdown("**📎 Upload Files**")
    sidebar_upload = st.file_uploader(
        "Choose files",
        type=['txt', 'pdf', 'docx', 'doc', 'png', 'jpg', 'jpeg', 'gif', 'bmp'],
        key="sidebar_uploader",
        label_visibility="collapsed",
        help="Upload documents or images for context (PDF, DOCX, TXT, Images with OCR)"
    )
    if sidebar_upload and sidebar_upload.name not in st.session_state.document_contents:
        file_details = {
            "filename": sidebar_upload.name,
            "filetype": sidebar_upload.type,
            "filesize": sidebar_upload.size
        }
        with st.spinner(f"Processing {sidebar_upload.name}..."):
            st.session_state.document_contents[sidebar_upload.name] = process_document(sidebar_upload)
        st.session_state.uploaded_files.append(file_details)
        st.sidebar.success(f"✅ {sidebar_upload.name} uploaded!")
        st.rerun()

    # Sidebar tips (Compact)
    st.markdown("---")
    st.markdown("**💡 Quick Tips:**")
    st.markdown("- Ask about studies or homework")
    st.markdown("- Upload files for context")
    st.markdown("- Be specific for better responses")

    # Show uploaded files (Styled)
    if st.session_state.uploaded_files:
        st.markdown("---")
        st.markdown("**📁 Your Files:**")
        for i, f in enumerate(st.session_state.uploaded_files):
            col1, col2 = st.columns([3, 1])
            with col1:
                st.markdown(f"📄 {f['filename']}")
            with col2:
                if st.button("🗑️", key=f"del_file_{i}", help="Remove file"):
                    fname = f['filename']
                    st.session_state.uploaded_files.pop(i)
                    st.session_state.document_contents.pop(fname, None)
                    st.rerun()

# ---------------------------
# Header (Exact Match to Quiz Structure: LectureBuddies - Chatbot)
# ---------------------------
st.markdown("<h1 class='main-title'>LectureBuddies</h1>", unsafe_allow_html=True)
st.markdown("<h2 style='text-align: center; color: #4e54c8; font-weight: 600; font-family: Poppins, sans-serif;'>Lecturebuddies - Chatbot</h2>", unsafe_allow_html=True)
st.markdown("<p class='tagline'>Your intelligent study companion—chat, summarize, and learn with AI ✨</p>", unsafe_allow_html=True)
st.markdown("<hr class='gradient'>", unsafe_allow_html=True)

# ---------------------------
# Quick Actions (Styled to Match)
# ---------------------------
if not st.session_state.messages:
    st.markdown(
        """
        <div class="welcome-container">
            <h3 class="welcome-title">👋 Welcome to LectureBuddies Chat!</h3>
            <p class="welcome-text">Start chatting or try a quick action below to dive into your studies.</p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    col1, col2, col3 = st.columns(3)
    presets = {
        "📚 Help with Homework": "I need help understanding this homework assignment. Can you explain step-by-step?",
        "🔬 Explain a Concept": "I'm studying this concept but finding it difficult. Can you explain clearly with examples?",
        "💡 Study Tips": "I want to improve my study efficiency. What study strategies should I use?"
    }
    for col, (label, prompt) in zip([col1, col2, col3], presets.items()):
        with col:
            if st.button(label, key=label.replace(" ", "_").lower(), help="Start with this prompt"):
                st.session_state.messages.append({"role": "user", "content": prompt})
                with st.spinner("🤔 Thinking..."):
                    reply = get_groq_response(prompt)
                st.session_state.messages.append({"role": "assistant", "content": reply})
                st.rerun()

# ---------------------------
# Chat Display (Direct, No Container)
# ---------------------------
if st.session_state.messages:
    for msg in st.session_state.messages:
        if msg["role"] == "user":
            st.markdown(f"""
            <div class="user-message">
                <div class="user-bubble">
                    {msg['content']}
                </div>
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown(f"""
            <div class="assistant-message">
                <div class="assistant-bubble">
                    {msg['content']}
                </div>
            </div>
            """, unsafe_allow_html=True)

# ---------------------------
# Chat Input (Professional Styling, Shorter Placeholder)
# ---------------------------
user_input = st.chat_input(placeholder="💬 Ask about studies or uploaded files...")
if user_input and user_input.strip():
    # 🔑 Inject file content here
    final_input = inject_file_content(user_input.strip())

    st.session_state.messages.append({"role": "user", "content": user_input.strip()})
    with st.spinner("🤔 Thinking..."):
        reply = get_groq_response(final_input)
    st.session_state.messages.append({"role": "assistant", "content": reply})
    st.rerun()

# ---------------------------
# Footer (Compact)
# ---------------------------
st.markdown("---")
st.markdown(
    "<p style='text-align: center; color: #888; font-size: 12px; font-family: Poppins, sans-serif; margin-bottom: 0;'>"
    "© 2023 LectureBuddies | Built with ❤️ for educational excellence | Powered by Groq AI</p>",
    unsafe_allow_html=True
)
